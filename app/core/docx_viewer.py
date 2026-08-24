"""Navegação e anotação em documentos DOCX."""

import html
import re
import xml.etree.ElementTree as ET
from datetime import datetime, timezone
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from loguru import logger
from rapidfuzz import fuzz

from app.models.schemas import TextLocation

_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_REL_NS = "{http://schemas.openxmlformats.org/package/2006/relationships}"
_CT_NS = "{http://schemas.openxmlformats.org/package/2006/content-types}"
_COMMENT_REL_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
)
_COMMENT_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
)
_COMMENT_AUTHOR = "BGF Revisão"
_COMMENT_INITIALS = "BGF"


def find_text_in_docx(file_path: str, query: str, max_results: int = 5) -> list[TextLocation]:
    """Localiza trecho nos parágrafos do DOCX (mesmo índice do XML, inclui tabelas)."""
    query = re.sub(r"\s+", " ", (query or "").strip())
    if len(query) < 4:
        return []

    path = Path(file_path)
    if not path.exists() or path.suffix.lower() not in (".docx", ".doc"):
        return []

    locations: list[TextLocation] = []
    for idx, text in enumerate(iter_docx_paragraph_texts(file_path)):
        text = (text or "").strip()
        if not text:
            continue
        score = 100 if query in text else fuzz.partial_ratio(query, text)
        if score >= 75:
            locations.append(
                TextLocation(
                    page=idx + 1,
                    paragraph_index=idx,
                    text=text[:500],
                    match_score=score / 100.0,
                    document_type="docx",
                )
            )
            if len(locations) >= max_results:
                break

    return locations


def render_docx_paragraphs_html(
    file_path: str,
    highlight_indices: set[int] | None = None,
    highlight_color: str = "#fff3cd",
) -> str:
    """Renderiza parágrafos do DOCX como HTML com destaques."""
    doc = Document(file_path)
    parts = [
        "<style>.docx-view{font-family:Georgia,serif;line-height:1.7;padding:16px;"
        "background:#fafafa;border-radius:8px;max-height:70vh;overflow-y:auto}"
        ".docx-p{margin:8px 0;padding:8px;border-radius:4px}"
        ".docx-hl{background:" + highlight_color + ";border-left:4px solid #f59e0b}"
        ".docx-num{color:#6b7280;font-size:12px}</style><div class='docx-view'>",
    ]
    highlight_indices = highlight_indices or set()
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        safe = html.escape(text)
        cls = "docx-p docx-hl" if idx in highlight_indices else "docx-p"
        parts.append(f'<p class="{cls}"><span class="docx-num">¶{idx+1}</span> {safe}</p>')
    parts.append("</div>")
    return "".join(parts)


def add_comment_to_docx(
    file_path: str,
    paragraph_index: int,
    comment_text: str,
    output_path: str | None = None,
) -> str:
    """Insere comentário nativo do Word (balão) no parágrafo indicado."""
    return _add_native_word_comment(
        file_path,
        paragraph_index,
        comment_text,
        output_path=output_path,
    )


def _para_plain_text(p_el: ET.Element) -> str:
    parts: list[str] = []
    for node in p_el.iter(f"{_W_NS}t"):
        if node.text:
            parts.append(node.text)
        if node.tail:
            parts.append(node.tail)
    return re.sub(r"\s+", " ", "".join(parts)).strip()


def _iter_body_paragraphs(body: ET.Element):
    """Parágrafos do corpo em ordem de documento (inclui células de tabela)."""

    def walk_tbl(tbl: ET.Element):
        for row in tbl.findall(f"./{_W_NS}tr"):
            for cell in row.findall(f"./{_W_NS}tc"):
                for el in list(cell):
                    if el.tag == f"{_W_NS}p":
                        yield el
                    elif el.tag == f"{_W_NS}tbl":
                        yield from walk_tbl(el)

    for child in list(body):
        if child.tag == f"{_W_NS}p":
            yield child
        elif child.tag == f"{_W_NS}tbl":
            yield from walk_tbl(child)

def iter_docx_paragraph_texts(file_path: str) -> list[str]:
    """Lista de textos de parágrafo na mesma ordem da extração de comentários."""
    path = Path(file_path)
    if not path.exists():
        return []
    try:
        with ZipFile(file_path) as zf:
            if "word/document.xml" not in zf.namelist():
                return []
            doc_root = ET.fromstring(zf.read("word/document.xml"))
            body = doc_root.find(f".//{_W_NS}body")
            if body is None:
                return []
            return [_para_plain_text(p) for p in _iter_body_paragraphs(body)]
    except Exception:
        doc = Document(file_path)
        return [p.text.strip() for p in doc.paragraphs]


def _load_comment_meta(zf: ZipFile) -> dict[str, dict]:
    if "word/comments.xml" not in zf.namelist():
        return {}
    root = ET.fromstring(zf.read("word/comments.xml"))
    meta: dict[str, dict] = {}
    for cmt_el in root.findall(f".//{_W_NS}comment"):
        cid = cmt_el.get(f"{_W_NS}id")
        if cid is None:
            continue
        author = cmt_el.get(f"{_W_NS}author") or "Desconhecido"
        parts: list[str] = []
        for t in cmt_el.iter(f"{_W_NS}t"):
            if t.text:
                parts.append(t.text)
            if t.tail:
                parts.append(t.tail)
        meta[cid] = {"text": "".join(parts).strip(), "author": author}
    return meta


def _extract_docx_comments_from_xml(file_path: str) -> list[dict]:
    """Extrai comentários com âncora real (commentRangeStart/End), inclusive em tabelas."""
    from app.core.extractor import compute_comment_stable_id

    comments: list[dict] = []
    try:
        with ZipFile(file_path) as zf:
            if "word/document.xml" not in zf.namelist():
                return []
            doc_root = ET.fromstring(zf.read("word/document.xml"))
            comment_meta = _load_comment_meta(zf)
            body = doc_root.find(f".//{_W_NS}body")
            if body is None:
                return []

            para_idx = -1
            active: dict[str, dict] = {}
            closed_ids: set[str] = set()

            for child in _iter_body_paragraphs(body):
                para_idx += 1
                para_text = _para_plain_text(child)

                # Starts e ends no mesmo parágrafo: processar em ordem do XML
                for node in child.iter():
                    if node.tag == f"{_W_NS}commentRangeStart":
                        cid = node.get(f"{_W_NS}id")
                        if cid and cid not in closed_ids:
                            active[cid] = {"start_para": para_idx, "texts": []}
                    elif node.tag == f"{_W_NS}commentRangeEnd":
                        cid = node.get(f"{_W_NS}id")
                        if cid not in active:
                            continue
                        data = active.pop(cid)
                        if para_text and para_idx >= data["start_para"]:
                            # garante o parágrafo do end na âncora
                            if not data["texts"] or data["texts"][-1] != para_text:
                                data["texts"].append(para_text)
                        referenced = " ".join(data["texts"]).strip()
                        meta = comment_meta.get(cid, {})
                        text = meta.get("text", "")
                        author = meta.get("author", "Desconhecido")
                        start_para = data["start_para"]
                        stable_id = compute_comment_stable_id(
                            file_path,
                            page=start_para + 1,
                            author=author,
                            comment_text=text,
                            referenced_text=referenced,
                            paragraph_index=start_para,
                        )
                        comments.append(
                            {
                                "id": stable_id,
                                "stable_id": stable_id,
                                "page": start_para + 1,
                                "comment_text": text,
                                "referenced_text": referenced,
                                "author": author,
                                "date": "",
                                "paragraph_index": start_para,
                                "word_comment_id": cid,
                            }
                        )
                        closed_ids.add(cid)

                # Acumula texto nos ranges abertos (após processar starts deste parágrafo)
                for data in active.values():
                    if para_text and para_idx >= data["start_para"]:
                        if not data["texts"] or data["texts"][-1] != para_text:
                            data["texts"].append(para_text)

            # Comentários com range open (malformado): ainda inclui pelo meta
            for cid, data in list(active.items()):
                if cid in closed_ids:
                    continue
                meta = comment_meta.get(cid, {})
                text = meta.get("text", "")
                if not text.strip():
                    continue
                author = meta.get("author", "Desconhecido")
                start_para = data["start_para"]
                referenced = " ".join(data["texts"]).strip()
                stable_id = compute_comment_stable_id(
                    file_path,
                    page=start_para + 1,
                    author=author,
                    comment_text=text,
                    referenced_text=referenced,
                    paragraph_index=start_para,
                )
                comments.append(
                    {
                        "id": stable_id,
                        "stable_id": stable_id,
                        "page": start_para + 1,
                        "comment_text": text,
                        "referenced_text": referenced,
                        "author": author,
                        "date": "",
                        "paragraph_index": start_para,
                        "word_comment_id": cid,
                    }
                )

            # Garante 100% de cobertura: meta em comments.xml sem âncora visitada
            seen_ids = {c.get("word_comment_id") for c in comments}
            for cid, meta in comment_meta.items():
                if cid in seen_ids:
                    continue
                text = meta.get("text", "")
                if not text.strip():
                    continue
                author = meta.get("author", "Desconhecido")
                stable_id = compute_comment_stable_id(
                    file_path,
                    page=1,
                    author=author,
                    comment_text=text,
                    paragraph_index=None,
                )
                comments.append(
                    {
                        "id": stable_id,
                        "stable_id": stable_id,
                        "page": 1,
                        "comment_text": text,
                        "referenced_text": "",
                        "author": author,
                        "date": "",
                        "paragraph_index": None,
                        "word_comment_id": cid,
                    }
                )
    except Exception as exc:
        logger.debug("XML de comentários DOCX: {}", exc)
    return comments


def extract_comments_from_docx(file_path: str) -> list[dict]:
    """Extrai comentários nativos do DOCX com posição no parágrafo ancorado."""
    from app.core.extractor import compute_comment_stable_id

    comments = _extract_docx_comments_from_xml(file_path)
    if comments:
        logger.info("Extraídos {} comentário(s) DOCX (XML) de {}", len(comments), Path(file_path).name)
        return comments

    try:
        doc = Document(file_path)
        if hasattr(doc, "comments") and doc.comments:
            for i, cmt in enumerate(doc.comments):
                text = getattr(cmt, "text", str(cmt)) or ""
                author = getattr(cmt, "author", "Desconhecido") or "Desconhecido"
                stable_id = compute_comment_stable_id(
                    file_path,
                    page=i + 1,
                    author=author,
                    comment_text=text,
                )
                comments.append(
                    {
                        "id": stable_id,
                        "stable_id": stable_id,
                        "page": 1,
                        "comment_text": text,
                        "referenced_text": "",
                        "author": author,
                        "date": "",
                        "paragraph_index": None,
                    }
                )
    except Exception as exc:
        logger.debug("Comentários DOCX não disponíveis: {}", exc)
    return comments


def _xml_register_word_ns() -> None:
    ET.register_namespace("w", _W)


def _serialize_word_xml(root: ET.Element) -> bytes:
    _xml_register_word_ns()
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _ensure_comments_part(files: dict[str, bytes]) -> tuple[ET.Element, int]:
    """Garante comments.xml + content types + relacionamento. Retorna (root, próximo id)."""
    _xml_register_word_ns()
    if "word/comments.xml" in files and files["word/comments.xml"].strip():
        root = ET.fromstring(files["word/comments.xml"])
    else:
        root = ET.fromstring(
            b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            b'<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )

    max_id = -1
    for cmt in root.findall(f"{_W_NS}comment"):
        try:
            max_id = max(max_id, int(cmt.get(f"{_W_NS}id") or "0"))
        except ValueError:
            continue

    ct_xml = files["[Content_Types].xml"].decode("utf-8")
    if "/word/comments.xml" not in ct_xml:
        files["[Content_Types].xml"] = ct_xml.replace(
            "</Types>",
            '<Override PartName="/word/comments.xml" '
            f'ContentType="{_COMMENT_CONTENT_TYPE}"/></Types>',
            1,
        ).encode("utf-8")

    rels_path = "word/_rels/document.xml.rels"
    rels_xml = files[rels_path].decode("utf-8")
    if _COMMENT_REL_TYPE not in rels_xml:
        rid = _next_rel_id_from_xml(rels_xml)
        files[rels_path] = rels_xml.replace(
            "</Relationships>",
            f'<Relationship Id="{rid}" Type="{_COMMENT_REL_TYPE}" Target="comments.xml"/>'
            "</Relationships>",
            1,
        ).encode("utf-8")

    return root, max_id + 1


def _next_rel_id_from_xml(rels_xml: str) -> str:
    used = {int(n) for n in re.findall(r'Id="rId(\d+)"', rels_xml)}
    n = 1
    while n in used:
        n += 1
    return f"rId{n}"


def _append_comment_xml(root: ET.Element, comment_id: int, comment_text: str) -> None:
    now = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    cmt = ET.SubElement(
        root,
        f"{_W_NS}comment",
        {
            f"{_W_NS}id": str(comment_id),
            f"{_W_NS}author": _COMMENT_AUTHOR,
            f"{_W_NS}date": now,
            f"{_W_NS}initials": _COMMENT_INITIALS,
        },
    )
    p_el = ET.SubElement(cmt, f"{_W_NS}p")
    run = ET.SubElement(p_el, f"{_W_NS}r")
    t_el = ET.SubElement(run, f"{_W_NS}t")
    t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t_el.text = comment_text


def _anchor_paragraph(para: ET.Element, comment_id: int) -> None:
    """Envolve o parágrafo com commentRangeStart/End + commentReference."""
    start = ET.Element(f"{_W_NS}commentRangeStart", {f"{_W_NS}id": str(comment_id)})
    end = ET.Element(f"{_W_NS}commentRangeEnd", {f"{_W_NS}id": str(comment_id)})
    ref_run = ET.Element(f"{_W_NS}r")
    ET.SubElement(ref_run, f"{_W_NS}commentReference", {f"{_W_NS}id": str(comment_id)})

    children = list(para)
    insert_at = 0
    if children and children[0].tag == f"{_W_NS}pPr":
        insert_at = 1
    para.insert(insert_at, start)
    para.append(end)
    para.append(ref_run)


def _add_native_word_comment(
    file_path: str,
    paragraph_index: int,
    comment_text: str,
    output_path: str | None = None,
) -> str:
    from app.core.annotation_io import (
        AnnotationError,
        assert_writable,
        atomic_replace,
        make_sibling_temp,
        validate_comment_text,
        validate_document_path,
    )

    comment_text = validate_comment_text(comment_text)
    src = validate_document_path(file_path)
    dest = Path(output_path).resolve() if output_path else src
    if dest.suffix.lower() not in (".docx", ".doc"):
        dest = dest.with_suffix(".docx")
    if dest != src:
        import shutil

        dest.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(src, dest)
    assert_writable(dest)

    with ZipFile(dest, "r") as zf:
        files = {name: zf.read(name) for name in zf.namelist()}

    if "word/document.xml" not in files:
        raise AnnotationError("DOCX inválido: falta word/document.xml.")

    comments_root, new_id = _ensure_comments_part(files)
    doc_root = ET.fromstring(files["word/document.xml"])
    body = doc_root.find(f"{_W_NS}body")
    if body is None:
        body = doc_root.find(f".//{_W_NS}body")
    if body is None:
        raise AnnotationError("DOCX inválido: documento sem corpo.")

    paragraphs = list(_iter_body_paragraphs(body))
    if not paragraphs:
        raise AnnotationError("Não há parágrafos no DOCX para ancorar o comentário.")
    idx = int(paragraph_index)
    if idx < 0 or idx >= len(paragraphs):
        raise AnnotationError(
            f"Parágrafo {idx + 1} inválido — o documento tem {len(paragraphs)} parágrafo(s)."
        )

    _append_comment_xml(comments_root, new_id, comment_text)
    _anchor_paragraph(paragraphs[idx], new_id)

    files["word/comments.xml"] = _serialize_word_xml(comments_root)
    files["word/document.xml"] = _serialize_word_xml(doc_root)

    tmp = make_sibling_temp(dest)
    try:
        with ZipFile(tmp, "w", compression=ZIP_DEFLATED) as zf:
            for name, data in files.items():
                zf.writestr(name, data)
        atomic_replace(tmp, dest)
    except Exception:
        if tmp.exists():
            tmp.unlink(missing_ok=True)
        raise

    logger.info("Comentário nativo DOCX inserido no parágrafo {} → {}", idx, dest.name)
    return str(dest)

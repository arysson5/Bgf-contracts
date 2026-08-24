"""Gravação de comentários no arquivo (PDF nativo e DOCX Word)."""

from __future__ import annotations

from pathlib import Path

import fitz
import pytest
from docx import Document

from app.core.annotation_io import AnnotationError, validate_comment_text
from app.core.docx_viewer import add_comment_to_docx, extract_comments_from_docx
from app.core.extractor import extract_comments_from_pdf
from app.core.pdf_viewer import add_comment_to_pdf
from app.utils.export_ui import _enforce_source_suffix


def test_empty_comment_is_rejected() -> None:
    with pytest.raises(AnnotationError):
        validate_comment_text("   ")


def test_export_keeps_source_extension() -> None:
    assert _enforce_source_suffix("contrato.txt", "a.pdf") == "contrato.pdf"
    assert _enforce_source_suffix("revisao", r"C:\docs\arquivo.docx") == "revisao.docx"
    assert _enforce_source_suffix("ok.pdf", "x.pdf") == "ok.pdf"


def test_pdf_comment_lands_on_click_point(tmp_path: Path) -> None:
    src = tmp_path / "alvo.pdf"
    doc = fitz.open()
    page = doc.new_page(width=400, height=600)
    page.insert_text((72, 200), "Clausula de pagamento em 30 dias.")
    doc.save(str(src))
    doc.close()

    add_comment_to_pdf(
        str(src),
        1,
        "Ajustar prazo de pagamento.",
        output_path=str(src),
        percent_point=(20.0, 35.0),
    )
    comments = extract_comments_from_pdf(str(src))
    assert comments, "Comentário não foi gravado no PDF"
    assert any("Ajustar prazo" in (c.get("comment_text") or "") for c in comments)
    rect = comments[0]["rects"][0]
    assert rect["y0"] > 50, "Comentário não deveria cair no canto superior (72,72)"


def test_pdf_comment_requires_location_when_search_fails(tmp_path: Path) -> None:
    src = tmp_path / "sem_ancora.pdf"
    doc = fitz.open()
    doc.new_page()
    doc.save(str(src))
    doc.close()
    with pytest.raises(AnnotationError):
        add_comment_to_pdf(str(src), 1, "Nota sem posição", output_path=str(src))


def test_docx_native_comment_roundtrip(tmp_path: Path) -> None:
    src = tmp_path / "alvo.docx"
    doc = Document()
    doc.add_paragraph("Primeiro parágrafo do contrato.")
    doc.add_paragraph("Segunda cláusula sobre o prazo.")
    doc.save(str(src))

    add_comment_to_docx(str(src), 1, "Favor alinhar o prazo com a proposta.", output_path=str(src))
    comments = extract_comments_from_docx(str(src))
    assert len(comments) >= 1
    assert any("alinhar o prazo" in (c.get("comment_text") or "") for c in comments)
    hit = next(c for c in comments if "alinhar o prazo" in (c.get("comment_text") or ""))
    assert hit.get("paragraph_index") == 1
    assert hit.get("author") == "BGF Revisão"

    # Texto do corpo não deve virar parágrafo amarelo
    body = Document(str(src))
    texts = [p.text for p in body.paragraphs]
    assert not any("[Contract Analyzer]" in t for t in texts)
